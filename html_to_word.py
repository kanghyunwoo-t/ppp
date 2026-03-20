from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

class HtmlToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self._set_korean_font()

    def _set_korean_font(self):
        """한국어 폰트('맑은 고딕') 기본 설정"""
        style = self.doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def parse_and_convert(self, html_content: str, output_path: str):
        """HTML을 파싱하여 Word 문서로 변환 및 저장합니다."""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 최상위 요소들을 순회하며 Word 객체로 매핑
        for element in soup.children:
            self._map_element_to_docx(element)
            
        self.doc.save(output_path)
        print(f"성공적으로 저장되었습니다: {output_path}")

    def _map_element_to_docx(self, element):
        """HTML 태그별 문서 매핑 로직"""
        # 1. 태그 없이 텍스트만 덩그러니 있는 경우 (놓치지 않고 문단으로 추가)
        if element.name is None:
            text = str(element).strip()
            if text:
                self.doc.add_paragraph(text)
            return
            
        tag = element.name.lower()
        
        # 2. 페이지 나누기
        if tag == 'hr':
            self.doc.add_page_break()
            
        # 3. 제목 태그 (h1~h6 모두 지원, 워드 한계인 4까지만 매핑)
        elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag.replace('h', ''))
            level = min(level, 4)
            self.doc.add_heading(element.get_text(strip=True), level=level)
            
        # 4. 문단 태그
        elif tag == 'p':
            self.doc.add_paragraph(element.get_text(strip=True))
            
        elif tag == 'ul' or tag == 'ol':
            for li in element.find_all('li', recursive=False):
                self.doc.add_paragraph(li.get_text(strip=True), style='List Bullet')
                
        # 6. 표 태그
        elif tag == 'table':
            self._process_table(element)
            
        # 7. 그 외 모든 알 수 없는 태그(div, span, b, figure 등)는 무시하지 않고 안으로 파고듦
        else:
            if hasattr(element, 'children'):
                for child in element.children:
                    self._map_element_to_docx(child)

    def _process_table(self, table_soup):
        """
        [고도화된 핵심 로직] HTML Table -> Word Table 매핑
        2차원 매트릭스 알고리즘을 사용하여 가로(colspan) 및 세로(rowspan) 병합을 완벽하게 처리합니다.
        """
        rows = table_soup.find_all('tr')
        if not rows:
            return

        num_rows = len(rows)
        max_cols = 0
        grid = {}  # (row_idx, col_idx) 좌표에 셀 정보를 저장할 딕셔너리

        # 1. HTML 표 구조를 2차원 그리드(Grid) 매트릭스로 분석
        for r_idx, row in enumerate(rows):
            c_idx = 0
            for cell in row.find_all(['td', 'th']):
                # 현재 행에서 이미 이전 줄의 세로 병합(rowspan)으로 차지된 셀은 건너뛰기
                while (r_idx, c_idx) in grid:
                    c_idx += 1
                
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))
                
                # 병합된 영역만큼 그리드 좌표에 정보 기록
                for r in range(rowspan):
                    for c in range(colspan):
                        grid[(r_idx + r, c_idx + c)] = {
                            'text': cell.get_text(strip=True) if r == 0 and c == 0 else '',
                            'is_primary': (r == 0 and c == 0),
                            'rowspan': rowspan,
                            'colspan': colspan
                        }
                c_idx += colspan
            max_cols = max(max_cols, c_idx)

        if max_cols == 0:
            return

        # 2. 파악된 최대 행/열 개수로 docx 표 생성
        docx_table = self.doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Table Grid'

        # 3. docx 표에 데이터 채우기 및 실제 셀 병합(merge) 수행
        for r_idx in range(num_rows):
            for c_idx in range(max_cols):
                cell_info = grid.get((r_idx, c_idx))
                
                # 데이터가 존재하고, 병합의 기준점(Primary)인 경우에만 텍스트 입력 및 병합 수행
                if cell_info and cell_info['is_primary']:
                    rs = cell_info['rowspan']
                    cs = cell_info['colspan']
                    text = cell_info['text']
                    
                    docx_cell = docx_table.cell(r_idx, c_idx)
                    docx_cell.text = text
                    
                    # 세로(rowspan) 또는 가로(colspan) 병합이 필요한 경우
                    if rs > 1 or cs > 1:
                        # docx 표의 범위를 벗어나지 않도록 안전장치 적용
                        end_r = min(r_idx + rs - 1, num_rows - 1)
                        end_c = min(c_idx + cs - 1, max_cols - 1)
                        
                        if end_r > r_idx or end_c > c_idx:
                            target_cell = docx_table.cell(end_r, end_c)
                            docx_cell.merge(target_cell)
